from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any, Tuple
from difflib import SequenceMatcher
import uvicorn
import io
import pdfplumber
import docx
import re
import json
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import os
try:
    from pdf2image import convert_from_bytes
    _PDF2IMAGE_AVAILABLE = True
except Exception:
    _PDF2IMAGE_AVAILABLE = False

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

app = FastAPI(title="Internship Matcher MVP")

# Allow CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.lower().endswith('.pdf'):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return '\n'.join(page.extract_text() or '' for page in pdf.pages)
    elif filename.lower().endswith('.docx'):
        doc = docx.Document(io.BytesIO(file_bytes))
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        return file_bytes.decode(errors='ignore')

def extract_text_and_links(file_bytes: bytes, filename: str) -> Tuple[str, list]:
    links = []
    text = ""
    if filename.lower().endswith('.pdf'):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or '') + '\n'
                # Robust hyperlink extraction
                try:
                    if hasattr(page, 'hyperlinks') and page.hyperlinks:
                        for link in page.hyperlinks:
                            if isinstance(link, dict) and 'uri' in link:
                                links.append(link['uri'])
                            elif isinstance(link, str):
                                links.append(link)
                    elif hasattr(page, 'annots') and page.annots:
                        for annot in page.annots or []:
                            uri = annot.get('uri')
                            if uri:
                                links.append(uri)
                except Exception as e:
                    pass  # Ignore hyperlink extraction errors
        # Fallback: image-based PDF -> OCR via pdf2image if little/no text
        if (not text or len(text.strip()) < 40) and _PDF2IMAGE_AVAILABLE:
            try:
                images = convert_from_bytes(file_bytes, dpi=200, fmt='png')
                ocr_pages = images[:3]  # limit for performance
                ocr_text = []
                for img in ocr_pages:
                    # Preprocess image for better OCR
                    def preprocess_image(img):
                        # Convert to grayscale
                        img = img.convert('L')
                        # Increase contrast
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(2)
                        # Apply a slight blur to reduce noise
                        img = img.filter(ImageFilter.MedianFilter())
                        return img

                    preprocessed_img = preprocess_image(img)

                    # Try multiple PSM modes for better OCR
                    psm_modes = [6, 11]  # Assume a single uniform block, then sparse text
                    img_ocr_text = ""
                    for psm in psm_modes:
                        config = f'--psm {psm} -c tessedit_char_whitelist=abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789@._-'
                        img_ocr_text = pytesseract.image_to_string(preprocessed_img, config=config)
                        # If email-like pattern found, break early
                        if re.search(r'[\w\.-]+@[\w\.-]+', img_ocr_text):
                            break

                    ocr_text.append(img_ocr_text)
                text = '\n'.join(ocr_text)
            except Exception:
                # if OCR fails, leave text as is
                pass
    elif filename.lower().endswith('.docx'):
        doc = docx.Document(io.BytesIO(file_bytes))
        text = '\n'.join([para.text for para in doc.paragraphs])
        # Extract hyperlinks from docx
        for para in doc.paragraphs:
            for run in para.runs:
                if hasattr(run, "hyperlink") and run.hyperlink:
                    links.append(run.hyperlink.target)
    elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
        image = Image.open(io.BytesIO(file_bytes))

        # Preprocess image for better OCR
        def preprocess_image(img):
            # Convert to grayscale
            img = img.convert('L')
            # Increase contrast
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2)
            # Apply a slight blur to reduce noise
            img = img.filter(ImageFilter.MedianFilter())
            return img

        preprocessed_img = preprocess_image(image)

        # Try multiple PSM modes for better OCR
        psm_modes = [6, 11]  # Assume a single uniform block, then sparse text
        ocr_text = ""
        for psm in psm_modes:
            config = f'--psm {psm} -c tessedit_char_whitelist=abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789@._-'
            ocr_text = pytesseract.image_to_string(preprocessed_img, config=config)
            # If email-like pattern found, break early
            if re.search(r'[\w\.-]+@[\w\.-]+', ocr_text):
                break

        text = ocr_text

    else:
        text = file_bytes.decode(errors='ignore')
    return text, links

def is_likely_name(s: str) -> bool:
    # Check if string looks like a person's name: 2-3 words, no special chars or digits, no emails or urls, avoid keywords
    if not s:   
        return False
    words = s.split()
    if len(words) < 2 or len(words) > 3:
        return False
    if re.search(r'[@\d]', s):
        return False
    if re.search(r'http|www|\.com|\.net|\.org', s, re.IGNORECASE):
        return False
    # Avoid common keywords
    if re.search(r'(curriculum|resume|cv|domain|email|phone|linkedin|github|degree|b\.tech|bachelor|master|engineering|computer|science)', s, re.IGNORECASE):
        return False
    # Require at least half words capitalized
    cap_words = sum(1 for w in words if w[0].isupper())
    if cap_words < len(words) // 2 + 1:
        return False
    return True

KNOWN_SKILLS = {
    # Languages
    "python","java","javascript","js","typescript","ts","c","c++","c#","go","golang","ruby","php","sql","html","html5","css","css3",
    # Frameworks/Libs
    "django","flask","fastapi","react","reactjs","node","nodejs","express","angular","vue","spring","spring boot","pandas","numpy","matplotlib","scikit-learn","sklearn",
    # Data/Tools
    "excel","power bi","powerbi","tableau","mongodb","mysql","postgresql","firebase","git","github","docker","kubernetes","aws","gcp","azure",
    # ML/DS
    "machine learning","ml","deep learning","dl","nlp","data analysis","data analytics","ai",
    # Additional from resume
    "dsa","dbms","oops concepts","web development","linux","streamlit","tailwind css","bootstrap","figma","oop","data visualisation","genai","prompt engineering","software engineering"
}

SKILL_SYNONYMS = {
    "js": "javascript",
    "ts": "typescript",
    "node": "nodejs",
    "spring boot": "spring",
    "py": "python",
    "html5": "html",
    "css3": "css",
    "ml": "machine learning",
    "dl": "deep learning",
    "reactjs": "react",
    "powerbi": "power bi",
    "sklearn": "scikit-learn"
}

def normalize_token(s: str) -> str:
    s = s.strip().lower()
    s = SKILL_SYNONYMS.get(s, s)
    return s

def extract_skills_from_text_fallback(text: str) -> List[str]:
    text_l = text.lower()
    found = set()
    # simple contains check for known skills
    for skill in KNOWN_SKILLS:
        if skill in text_l:
            found.add(normalize_token(skill))
    # split common delimited sequences (e.g., html/css/js)
    for token in re.split(r"[^a-zA-Z+#]+", text_l):
        token = token.strip()
        if not token:
            continue
        token = normalize_token(token)
        if token in KNOWN_SKILLS:
            found.add(token)
    return sorted(found)

def similar(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()

def skills_overlap_score(profile_skills: List[str], required_skills: List[str]) -> float:
    if not required_skills:
        return 0.0
    prof = [normalize_token(s) for s in profile_skills]
    req = [normalize_token(s) for s in required_skills]
    matched = 0
    for r in req:
        best = 0.0
        for p in prof:
            if p == r:
                best = 1.0
                break
            # substring or fuzzy closeness
            if p in r or r in p:
                best = max(best, 0.8)
            best = max(best, similar(p, r))
        if best >= 0.6:
            matched += 1
    return matched / len(req)

def extract_resume_info(text: str, links: list, filename: str | None = None) -> dict:
    def normalize_name_candidate(raw: str) -> str:
        # Remove non-letter separators except spaces and apostrophes
        cleaned = re.sub(r"[^A-Za-z'\s]", ' ', raw)
        cleaned = re.sub(r"\s+", ' ', cleaned).strip()
        # If ALL CAPS, convert to title case
        if cleaned and cleaned.upper() == cleaned:
            cleaned = cleaned.title()
        return cleaned
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = ''
    # Name extraction (case-insensitive, fallback to scanning first few lines for likely name)
    for line in lines:
        if re.search(r'^name\s*[:\-]', line, re.IGNORECASE):
            if ':' in line:
                name = line.split(':', 1)[1].strip()
            else:
                name = line.split('-', 1)[1].strip()
            break
    if not name and lines:
        # Scan first few lines (top matter) for a likely name
        for candidate in lines[:8]:
            candidate_norm = normalize_name_candidate(candidate)
            if is_likely_name(candidate_norm):
                name = candidate_norm
                break

    # Email extraction (text + hyperlinks)
    # Robust email regex
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    email = email_match.group(0) if email_match else ''
    # Fallback: search in lines for email-like patterns
    if not email:
        for line in lines:
            line_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', line)
            if line_match:
                email = line_match.group(0)
                break
    for link in links:
        if 'mailto:' in link.lower():
            email = link.split(':', 1)[-1]

    # Phone extraction (text only)
    phone_match = re.search(r'(\+?\d{1,3}[-\s]?)?\b\d{10}\b', text)
    phone = phone_match.group(0) if phone_match else ''

    # LinkedIn extraction (text + hyperlinks)
    linkedin_match = re.search(r'(https?://[^\s]+linkedin[^\s]+)', text, re.IGNORECASE)
    linkedin = linkedin_match.group(0) if linkedin_match else ''
    for link in links:
        if 'linkedin' in link.lower():
            linkedin = link

    # GitHub extraction (text + hyperlinks)
    github_match = re.search(r'(https?://[^\s]+github[^\s]+)', text, re.IGNORECASE)
    github = github_match.group(0) if github_match else ''
    for link in links:
        if 'github' in link.lower():
            github = link

    # Name fallbacks from email / LinkedIn / filename if still missing
    if not name:
        if email:
            local = email.split('@', 1)[0]
            parts = re.split(r'[._\-]+', local)
            parts = [p for p in parts if p and not p.isdigit()]
            if len(parts) >= 2:
                candidate = normalize_name_candidate(' '.join(parts[:3]))
                if is_likely_name(candidate):
                    name = candidate
        if not name and linkedin:
            # Try to parse /in/slug or last path segment
            m = re.search(r'/in/([^/?#]+)', linkedin)
            slug = m.group(1) if m else linkedin.rstrip('/').split('/')[-1]
            slug = re.split(r'[?&#]', slug)[0]
            parts = re.split(r'[._\-]+', slug)
            parts = [p for p in parts if p and p.lower() not in ('in', 'profile', 'pub')]
            if len(parts) >= 2:
                candidate = normalize_name_candidate(' '.join(parts[:3]))
                if is_likely_name(candidate):
                    name = candidate
        if not name and filename:
            base = os.path.splitext(os.path.basename(filename))[0]
            parts = re.split(r'[._\-]+', base)
            parts = [p for p in parts if p and p.lower() not in ('cv', 'resume') and not p.isdigit()]
            if len(parts) >= 2:
                candidate = normalize_name_candidate(' '.join(parts[:3]))
                if is_likely_name(candidate):
                    name = candidate

    # Skills extraction
    skills = []
    # Look for TECHNICAL SKILLS section
    tech_skills_start = False
    for i, line in enumerate(lines):
        if re.search(r'technical skills', line, re.IGNORECASE):
            tech_skills_start = True
            continue
        if tech_skills_start:
            if line == '' or re.search(r'certification|project|experience|education', line, re.IGNORECASE):
                break
            # Extract from lines like "Languages:Python, C, C++"
            if ':' in line:
                skill_list = [normalize_token(s.strip()) for s in line.split(':')[-1].split(',') if s.strip()]
                skills.extend(skill_list)
    # Look for COURSEWORK / SKILLS section
    coursework_start = False
    for i, line in enumerate(lines):
        if re.search(r'coursework.*skills', line, re.IGNORECASE):
            coursework_start = True
            continue
        if coursework_start:
            if line == '' or re.search(r'project|education|technical skills|certification', line, re.IGNORECASE):
                break
            # Extract bullet points
            if line.startswith('•'):
                skill = normalize_token(line[1:].strip())
                if skill in KNOWN_SKILLS:
                    skills.append(skill)
    # Dedupe and sort
    skills = sorted(set(skills))
    if not skills:
        skills = extract_skills_from_text_fallback(text)

    # Experience extraction (only headline/role/company)
    experience = []
    exp_start = False
    for line in lines:
        if re.search(r'experience|work experience', line, re.IGNORECASE):
            exp_start = True
            continue
        if exp_start:
            # Stop if section changes
            if line == '' or re.search(r'project|education|skills|certification', line, re.IGNORECASE):
                break
            # Only pick lines that look like a headline (role/company)
            # Example: contains 'at', ',', '-', '|', or ends with 'Ltd', 'Inc', 'Corp', etc.
            if re.search(r'\b(at|,|-|\||Ltd|Inc|Corp|Pvt|LLC)\b', line):
                experience.append(line)

    # Projects extraction
    projects = []
    proj_start = False
    for line in lines:
        if re.search(r'project', line, re.IGNORECASE):
            proj_start = True
            continue
        if proj_start:
            if line == '' or re.search(r'experience|education|skills|certification', line, re.IGNORECASE):
                break
            projects.append(line)

    # Education extraction
    education = []
    edu_start = False
    for line in lines:
        if re.search(r'education', line, re.IGNORECASE):
            edu_start = True
            continue
        if edu_start:
            if line == '' or re.search(r'project|experience|skills|certification', line, re.IGNORECASE):
                break
            education.append(line)

    # Certifications extraction
    certifications = []
    cert_start = False
    for line in lines:
        if re.search(r'certification', line, re.IGNORECASE):
            cert_start = True
            continue
        if cert_start:
            if line == '' or re.search(r'project|experience|skills|education', line, re.IGNORECASE):
                break
            certifications.append(line)

    # Location extraction
    location = ''
    for line in lines:
        if re.search(r'location|city', line, re.IGNORECASE):
            location = line.split(':')[-1].strip()
            break

    # Preferences extraction
    domains = []
    locations = []
    for line in lines:
        if re.search(r'domain[:\-]', line, re.IGNORECASE):
            domains = [d.strip() for d in line.split(':')[-1].split(',')]
        if re.search(r'location[:\-]', line, re.IGNORECASE):
            locations = [l.strip() for l in line.split(':')[-1].split(',')]

    # Infer candidate role/domain from text if not provided
    inferred_domains = []
    text_l = text.lower()
    if any(k in text_l for k in ["data science","data scientist","data analyst","analytics"]):
        inferred_domains.append("Data Science")
    if any(k in text_l for k in ["backend","fastapi","django","flask","api development"]):
        inferred_domains.append("Software Development")
    if any(k in text_l for k in ["frontend","react","angular","vue","html","css","javascript"]):
        inferred_domains.append("Frontend")
    if any(k in text_l for k in ["fullstack","full-stack","mern","mean","spring","nodejs","react"]):
        inferred_domains.append("Fullstack")
    # Merge and dedupe
    if inferred_domains:
        domains = list({*domains, *inferred_domains})

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "location": location,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
        "certifications": certifications,
        "preferences": {
            "domains": domains,
            "locations": locations
        }
    }

def smart_extract(text: str, links: list, filename: str | None = None) -> Dict[str, Any]:
    info = extract_resume_info(text, links, filename)
    return info

def score_internship(profile: Dict[str, Any], post: Dict[str, Any]) -> float:
    score = 0.0
    profile_skills = profile.get("skills", []) or []
    required_skills = post.get("required_skills", []) or []
    skill_score = skills_overlap_score(profile_skills, required_skills)  # 0..1
    score += 0.7 * skill_score

    # Domain preference match
    domain_pref = [d.lower() for d in profile.get("preferences", {}).get("domains", [])]
    post_domain = (post.get("domain") or "").lower()
    if post_domain and any(similar(d, post_domain) >= 0.7 or d in post_domain or post_domain in d for d in domain_pref):
        score += 0.2

    # Location preference match
    location_pref = [l.lower() for l in profile.get("preferences", {}).get("locations", [])]
    post_location = (post.get("location") or "").lower()
    if post_location and any(similar(l, post_location) >= 0.7 or l in post_location or post_location in l for l in location_pref):
        score += 0.1

    return max(0.0, min(score, 1.0))

def load_postings() -> List[Dict[str, Any]]:
    # Check if the file exists before trying to open it
    if not os.path.isfile("internships.json"):
        return []  # Return an empty list or handle the error as needed
    with open("internships.json", "r", encoding="utf-8") as f:
        return json.load(f)

from datetime import datetime

def parse_date_safe(s: str) -> datetime | None:
    try:
        return datetime.fromisoformat(s)
    except Exception:
        try:
            return datetime.strptime(s, "%Y-%m-%d")
        except Exception:
            return None

def match_internships(profile: Dict[str, Any]) -> List[Dict[str, Any]]:
    postings = load_postings()
    results_all = []
    for p in postings:
        s = score_internship(profile, p)
        # Recency boost (up to +0.1)
        pd = p.get("posting_date")
        boost = 0.0
        if pd:
            dt = parse_date_safe(pd)
            if dt:
                days = max(0, (datetime.now() - dt).days)
                # 0 days => +0.1, 30+ days => ~0
                boost = max(0.0, 0.1 * (1 - min(days, 30) / 30))
        s = min(1.0, s + boost)
        results_all.append({**p, "score": s})
    # Sort all by score
    results_all.sort(key=lambda x: x["score"], reverse=True)
    # Primary filter: keep strong matches
    strong = [r for r in results_all if r["score"] > 0.3]
    # Fallback: if nothing strong, show top 3 closest anyway
    chosen = strong if strong else results_all[:3]
    return chosen[:5]

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    content = await file.read()
    text, links = extract_text_and_links(content, file.filename)
    profile = smart_extract(text, links, file.filename)
    results = match_internships(profile)
    return JSONResponse(content={"profile": profile, "internships": results})

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

